# -*- coding: utf-8 -*-
"""
Created on Tue Oct 22 16:12:25 2024

@author: huber
"""
from pathlib import Path
import pandas as pd
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

from config import *
from python_data_analysis_utils.utils.dataframe_utils import read_csv_file, filter_dataframe

#%%
csv_file_path = DATA_DIR / 'prepped' / 'final_feature_set.csv'
df_final = read_csv_file(csv_file_path, delimiter=',', low_memory=False, index_col=None, dtype={'DateY': str, 'DateM': str})

#%%

#================================================================================
# Table - Group
#================================================================================

# Oblicz liczności dla każdej grupy
group_counts = df_final['Group'].value_counts().sort_index()

# Tworzenie dokumentu Word
doc = Document()

# Dodanie numeru i tytułu tabeli w stylu APA
heading = doc.add_paragraph('Table 1\nDemographic Characteristics')
heading.alignment = WD_ALIGN_PARAGRAPH.LEFT

# Tworzenie tabeli (2 kolumny: Group, Frequency)
table = doc.add_table(rows=1, cols=2)
table.style = 'Table Grid'

# Dodanie nagłówków kolumn
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Group'
hdr_cells[1].text = 'Frequency'

# Wypełnienie tabeli danymi
for group, count in group_counts.items():
    row_cells = table.add_row().cells
    row_cells[0].text = str(group)
    row_cells[1].text = str(count)

# Zmiana rozmiaru czcionki na 12 (standard APA)
for row in table.rows:
    for cell in row.cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.size = Pt(12)

# Wyśrodkowanie nagłówków
for cell in table.rows[0].cells:
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

# Zapis dokumentu Word
doc.save(TABLES_DIR / 'Demographic_Characteristics.docx')

#%%


#================================================================================
# Table - Group / Fatal
#================================================================================
# Tworzenie nowej kolumny Group_Fatal łączącej wartości Group i Fatal
df_final['Group_Fatal'] = df_final['Group'] + df_final['Fatal'].astype(int).astype(str)

# Tworzenie nowej kolumny Group_Fatal łączącej wartości Group i Fatal
df_final['Group_Fatal'] = df_final['Group'] + df_final['Fatal'].astype(int).astype(str)

# Oblicz liczności dla każdej kombinacji Group i Fatal (Group_Fatal)
group_fatal_counts = df_final['Group_Fatal'].value_counts().sort_index()

# Tworzenie dokumentu Word
doc = Document()

# Dodanie numeru tabeli (pogrubiony) i tytułu (kursywa i Title Case)
heading = doc.add_paragraph()
heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
heading.add_run('Table 2').bold = True
#title = doc.add_paragraph()
title.add_run('Demographic Characteristics with Fatal Values').italic = True

# Tworzenie tabeli (2 kolumny: Group_Fatal, Frequency)
table = doc.add_table(rows=1, cols=2)
table.style = 'Table Grid'

# Dodanie nagłówków kolumn
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Group_Fatal'
hdr_cells[1].text = 'Frequency'

# Wypełnienie tabeli danymi
for group_fatal, count in group_fatal_counts.items():
    row_cells = table.add_row().cells
    row_cells[0].text = str(group_fatal)
    row_cells[1].text = f"{count:.0f}"  # Formatowanie liczb do 0 miejsc po przecinku

# Zmiana rozmiaru czcionki na 12 (standard APA)
for row in table.rows:
    for cell in row.cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.size = Pt(12)

# Wyśrodkowanie nagłówków
for cell in table.rows[0].cells:
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

# Dodanie notatki poniżej tabeli
note = doc.add_paragraph()
note.add_run("Note. 'Group_Fatal' represents the combination of group and fatal value (e.g., A0, A1, B0, B1).").italic = True

# Zapis dokumentu Word
doc.save(TABLES_DIR / 'Demographic_Characteristics_with_Fatal.docx')